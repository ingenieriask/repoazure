from os import name
from typing import Type
from django.core.mail.backends.smtp import EmailBackend
from django.core.mail import EmailMessage
from .utils_services import PDF
import logging
from datetime import datetime
from django.db import transaction, OperationalError
from django.db.models import Q
from django.conf import settings
import re
import io
import requests
import json
import pandas as pd
from datetime import date, timedelta
from enum import Enum
from django.core.exceptions import ValidationError
from core.models import AppParameter, ConsecutiveFormat, Consecutive, Country, FilingType, \
    Holiday, CalendarDay, CalendarDayType, Calendar, Notifications, SystemParameter, \
    SystemHelpParameter, Template, Task, ProceedingsConsecutiveFormat, ProceedingsConsecutive
from core.utils_services import FormatHelper
from django.contrib.auth.models import User, Permission
from correspondence.models import AlfrescoFile
from django.core.files import File
from django.core.files.temp import NamedTemporaryFile
from correspondence.ecm_services import ECMService
from docx import Document
import os
import threading
import time
from croniter import croniter
from django.apps import AppConfig
import pdfkit

logger = logging.getLogger(__name__)

class SystemParameterHelper():
    ''' '''
    
    @classmethod
    def get(cls, format_name):
        return SystemParameter.objects.get(name=format_name)

    @classmethod
    def get_json(cls, format_name):
        return json.loads(cls.get(format_name).value)


class SystemHelpParameterHelper():
    ''' '''
    
    @classmethod
    def get(cls, format_name):
        return SystemHelpParameter.objects.get(name=format_name)

    @classmethod
    def get_json(cls, format_name):
        return json.loads(cls.get(format_name).value)

class UserHelper():
    ''' '''
    
    @classmethod
    def list_by_permissions(cls, permission):
        return User.objects.filter(Q(groups__permissions__in=permission) | Q(user_permissions__in=permission)).distinct()
    
    @classmethod
    def list_by_permission_name(cls, permission_name):
        permission = Permission.objects.filter(codename=permission_name)
        return UserHelper.list_by_permissions(permission)

class Recipients():

    def __init__(self, email_to, email_cc=None, phone_to=None):

        self.email_to = [email_to] if isinstance(email_to, str) else email_to
        if email_cc:
            self.email_cc = [email_cc] if isinstance(email_cc, str) else email_cc
        else:
            self.email_cc = None
        if phone_to:
            self.phone_to = [phone_to] if isinstance(phone_to, str) else phone_to
        else:
            self.phone_to = None

class NotificationsHandler(object):
    '''Notification sender'''

    _params = {}
    sms_api_endpoint = 'https://api103.hablame.co/api/sms/v3/send/marketing'

    def get_params(func):
        '''Lazy load of email database parameters'''

        def wrapper(*args, **kwargs):
            # Lazy load
            if not NotificationsHandler._params:
                # Get only email related parameters
                qs = AppParameter.objects.filter(name__startswith='EMAIL_')
                NotificationsHandler._params = {
                    entry.name: entry.value for entry in qs}
            return func(*args, **kwargs)
        return wrapper

    _services = {
        'SEND_EMAIL': lambda data, email_format, recipients: NotificationsHandler.send_mail(
                    FormatHelper.replace_data(email_format.subject, data),
                    FormatHelper.replace_data(email_format.body, data),
                    NotificationsHandler._params['EMAIL_HOST_USER'],
                    recipients.email_to,
                    recipients.email_cc if recipients.email_cc else None
                ),
        'SEND_SMS': lambda data, email_format, recipients: NotificationsHandler.send_sms(
                    FormatHelper.replace_data(email_format.body_sms, data),
                    recipients.phone_to
                ),
    }

    @classmethod
    @get_params
    def send_notification(cls, format_name, data, recipients):
        try:
            email_format = Notifications.objects.get(name=format_name)

            for service in email_format.notifications_services.all():
                if service.name in cls._services:
                    if cls._services[service.name]:
                        cls._services[service.name](data, email_format, recipients)
                    else:
                        logger.error(f"service '{service}' unimplemented")
                else:
                    logger.error(f"service '{service}' undefined")

        except Exception as Error:
            print(f'Notification sending error: {Error}')
            logger.error(f'Notification sending error: {Error}')

    @classmethod
    @get_params
    def send_sms(cls, body='', to=None):
        '''Send SMS to several addresses using database parameters'''

        try:
            headers = SystemParameterHelper.get_json('SMS_PARAMETERS')
            if to:
                for numb in to:
                    data = {'toNumber':numb, 'sms': body}
                    r = requests.post(cls.sms_api_endpoint, headers=headers, json=data)
                    if r.ok:
                        print(json.loads(r.text))

        except Exception as Err:
            logger.error(Err) 

    @classmethod
    @get_params
    def send_mail(cls, subject='', body='', from_email=None, to=None, cc=None):
        '''Send emails to several addresses using database parameters'''

        if not from_email:
            from_email = cls._params['EMAIL_HOST_USER']

        # Configure email handler
        eb = EmailBackend(
            host=cls._params['EMAIL_HOST'],
            port=int(cls._params['EMAIL_PORT']),
            username=cls._params['EMAIL_HOST_USER'],
            password=cls._params['EMAIL_HOST_PASSWORD'],
            use_tls=bool(cls._params['EMAIL_USE_TLS']),
            fail_silently=False)
        try:
            eb.open()
            email = EmailMessage(subject, body, from_email, to, cc)
            email.content_subtype = "html"

            # Send email
            eb.send_messages([email])
            eb.close()
        except Exception as Error:
            logger.error('Error enviando el correo', Error)


class RecordCodeService(object):

    class Type(Enum):

        INPUT = 'input'
        OUTPUT = 'output'
        MEMO = 'memo'
        PROCERDINGS = 'proceedings'

    tokens = {
        'filing': ['{consecutive}', '{year}', '{type}'],
        'proceedings': ['{consecutive}', '{dependence}', '{subserie}', '{year}']
    }

    digits_token = 'consecutive'

    @classmethod
    def compile(cls, format, digits):
        '''Generate format from edited data'''

        code = re.sub(r'\s*', r'', format)
        code = re.sub(r',', r'', code)
        code = re.sub(f'{{{cls.digits_token}}}',
                      f'{{{cls.digits_token}:0{digits}d}}', code)
        return code

    @classmethod
    def decompile(cls, code):
        '''Decompile the format for edition'''

        if not code:
            return '', 10
        digits_pattern = r'\{' + cls.digits_token + r':0(\d+)d\}'
        code = re.sub(r'\s*', r'', code)
        format = code.replace('}', '},').replace('{', ',{').replace(
            ',,', ',').strip(',')
        format = re.sub(digits_pattern, f'{{{cls.digits_token}}}', format)
        r = re.search(digits_pattern, code)
        digits = int(r.group(1)) if r else 0
        return format, digits

    @classmethod
    @transaction.atomic
    def get_consecutive(cls, identifier):
        '''Retrieve the next consecutive code for a given type'''

        now = datetime.now()
        format = ConsecutiveFormat.objects.filter(
            effective_date__lte=now
        ).latest('effective_date').format

        filing_type = FilingType.objects.get(identifier=identifier.value)
        # Retrieve the last consecutive code
        try:
            consecutive = Consecutive.objects.get(type=filing_type)
        except Consecutive.DoesNotExist:
            consecutive = Consecutive(current=0, type=filing_type)

        # Update the consecutive code
        if consecutive.date.year != now.year:
            consecutive.current = 1
        else:
            consecutive.current += 1

        consecutive.date = now
        consecutive.save()

        # Generate formatted code
        params = {
            'type': filing_type.code,
            'year': now.year,
            'consecutive': consecutive.current
        }

        return format.format(**params)

    @classmethod
    @transaction.atomic
    def get_proceedings_consecutive(cls, subserie='02007', dependence='014'):
        '''Retrieve the next consecutive code for a given type'''

        now = datetime.now()
        format = ProceedingsConsecutiveFormat.objects.filter(
            effective_date__lte=now
        ).latest('effective_date').format

        # Retrieve the last consecutive code
        try:
            consecutive = ProceedingsConsecutive.objects.get()
        except ProceedingsConsecutive.DoesNotExist:
            consecutive = ProceedingsConsecutive(current=0)

        # Update the consecutive code
        if consecutive.date.year != now.year:
            consecutive.current = 1
        else:
            consecutive.current += 1

        consecutive.date = now
        consecutive.save()

        # Generate formatted code
        params = {
            'subserie': subserie, 
            'dependence': dependence,
            'year': now.year,
            'consecutive': consecutive.current
        }

        return format.format(**params)


class CalendarService(object):
    '''Service for calculate working days'''

    # https://date.nager.at/
    # API limitations:
    # 50 requests per day
    holiday_api_endpoint = 'https://date.nager.at/api/v3/publicholidays'

    @classmethod
    def get_days_of_year():
        return

    @classmethod
    def get_remaining_business_days(cls, request_day, max_response_days):
        return max_response_days - cls.get_business_days_since(request_day)

    @classmethod
    def get_business_days_since(cls, request_day):
        days = CalendarDay.objects.filter(
            date__range=[request_day, date.today()])
        workingdays = 0
        if days:
            for day in days:
                workingdays += 1 if day.type.name == "workingday" else 0
            return workingdays
        
        if isinstance(request_day, datetime):
            return (date.today() - request_day.date()).days
        return (date.today() - request_day).days

    @classmethod
    def get_calendar_days(cls, year):
        return CalendarDay.objects.filter(date__year=year)

    @classmethod
    def update_calendar_days(cls, year, nonworking_days):

        nonworking_days = {date(
            *map(int, d['id'].split('T')[0].split('-'))) for d in nonworking_days if 'id' in d}
        days = CalendarDay.objects.filter(date__year=year)
        working_day = CalendarDayType.objects.get(name='workingday')
        nonWorking_day = CalendarDayType.objects.get(name='non-workingday')
        nonWorking_day = CalendarDayType.objects.get(name='non-workingday')
        calendar = Calendar.objects.first()
        new = not days.exists()
        if new:
            day_range = pd.date_range(
                start=f'{year}-01-01',
                end=f'{int(year) + 1}-01-01',
                closed='left')
            days = [CalendarDay(date=d, calendar=calendar)
                    for d in day_range.date]
        for d in days:
            d.type = nonWorking_day if d.date in nonworking_days else working_day
        if new:
            return CalendarDay.objects.bulk_create(days)
        return CalendarDay.objects.bulk_update(days, ['type'])

    @classmethod
    def get_weekends(cls, year, week_day_code):
        '''Return the list of saturdays or sundays of an entire year'''

        return pd.date_range(
            start=f'{year}-01-01',
            end=f'{int(year) + 1}-01-01',
            freq=f'W-{week_day_code}',  # week_day[week_day_code],
            closed='left').tolist()

    @classmethod
    def get_holidays(cls, year, country_code):
        '''Return the list of holidays for a given year and country'''

        country_code = country_code.upper()

        holidays = Holiday.objects.filter(
            date__year=year,
            country__code=country_code,
        )
        if holidays.exists():
            return holidays
        # Retrieve data from external API if not exist in local database
        try:
            r = requests.get(
                f'{cls.holiday_api_endpoint}/{year}/{country_code}')
            if r.ok:
                json_response = json.loads(r.text)
                country = Country.objects.get(code=country_code)
                holidays = Holiday.objects.bulk_create([Holiday(**{
                    'date': datetime.strptime(h['date'], r'%Y-%m-%d'),
                    'country': country,
                    'local_name': h['localName']
                }) for h in json_response]
                )
                return holidays

        except Exception as Err:
            logger.error(Err)
        

class PdfCreationService(object):

    @classmethod
    def _save_pdf_in_ecm(cls, pdf, pqrs, name, extension):
        f = pdf.output(dest='S').encode('latin-1')
        document_temp_file = NamedTemporaryFile()
        document_temp_file.write(f)

        size = document_temp_file.tell()

        document_temp_file.seek(0)
        document_temp_file.flush()

        node_id = ECMService.upload(File(document_temp_file, name=name+extension), pqrs.folder_id)

        alfrescoFile = AlfrescoFile(cmis_id=node_id, 
                                    radicate=pqrs,
                                    name=name,
                                    extension=extension,
                                    size=int(size))
        alfrescoFile.save()
        return node_id

    @classmethod
    def create_pqrs_confirmation_label(cls, pqrs):
        
        # Create pdf instance        
        pdf = PDF()
        # Add page to pdf
        pdf.add_page()
        initial_x_pos = 30.0
        initial_y_pos = 50.0
        sizing_factor = 2
        border = 0
        # Create label base in header
        pdf.custom_header(pqrs, initial_x_pos, initial_y_pos, sizing_factor, border)
        # Save pdf file
        #pdf.output('label.pdf', 'F')
        PdfCreationService._save_pdf_in_ecm(pdf, pqrs, 'label', '.pdf')

    @classmethod
    def create_pqrs_summary(cls, pqrs):       
         
        # Create pdf instance    
        pdf = PDF()
        # Add page to pdf
        pdf.add_page()
        # Define left margin for the document
        pdf.set_left_margin(20.0)
        initial_x_pos = 100.0
        initial_y_pos = 12.0
        # Create custom doc header
        pdf.custom_header(pqrs, initial_x_pos, initial_y_pos, 1, 1)
        # Insert information inside doc body
        pdf.set_xy(20.0, initial_y_pos+60.0)
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0, 5, 'CIUDAD - '+pqrs.date_radicated.strftime('%Y-%m-%d')+'\nTRATAMIENTO\n')
        pdf.set_font('Arial', 'B', 12)
        pdf.multi_cell(0, 5, 'NOMBRES DESTINATARIO\n')
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0, 5, 'CARGO\nRAZ??N SOCIAl\nDIRECCI??N\nEMAIL\nMUNICIPIO - DEPARTAMENTO\n\n')
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(20, 5, 'ASUNTO: ')
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0,5, pqrs.subject + '\n\n')
        pdf.multi_cell(0,5, pqrs.data + '\n\n\n\n\n\nCordialmente,\n\n\n\nDATOS DE LA PERSONA QUE FIRMA\n')
        if pqrs.pqrsobject.principal_person.is_anonymous:
            pdf.multi_cell(0,5, 'AN??NIMO\n\n')
        else:
            pdf.multi_cell(0,5, pqrs.pqrsobject.principal_person.name+' '+pqrs.pqrsobject.principal_person.lasts_name+
                       '\nCARGO\n\n')
        pdf.set_font('Arial', '', 7)
        
        annexes = ''
        for index, val in enumerate(list(pqrs.files.all())):
            annexes += val.name
            if index != len(list(pqrs.files.all())) - 1:
                annexes += ', '
        
        pdf.multi_cell(0, 5, 'Anexo (s): ' + annexes)
        
        # Create custom doc footer
        pdf.custom_footer()
        # Save pdf file
        #pdf.output('summary.pdf', 'F')
        PdfCreationService._save_pdf_in_ecm(pdf, pqrs, pqrs.number, '.pdf')
    
    @classmethod
    def create_radicate_answer(cls, radicate, draft):
        # Create pdf instance    
        pdf = PDF()
        # Add page to pdf
        pdf.add_page()
        #add watermark if the document is a draft
        if draft:
            pdf.image('static/pqrs/img/draft_watermark.png', w=200, h=260)
        # Define left margin for the document
        pdf.set_left_margin(20.0)
        pdf.set_right_margin(-20.0)
        initial_x_pos = 100.0
        initial_y_pos = 12.0
        # Create custom doc header
        pdf.custom_header(radicate, initial_x_pos, initial_y_pos, 1, 1)
        # Insert information inside doc body
        pdf.set_xy(20.0, initial_y_pos+60.0)
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0, 5, 'CIUDAD - '+ radicate.date_radicated.strftime('%Y-%m-%d')+'\n\nSe??or(a)\n')
        pdf.set_font('Arial', 'B', 12)
        if radicate.person:
            if radicate.person.person_type.abbr == "PJ":
                pdf.multi_cell(0, 5, radicate.person.parent.representative+'\n')
            else:
                pdf.multi_cell(0, 5, radicate.person.name+' '+radicate.person.lasts_name+'\n')
            pdf.set_font('Arial', '', 12)
            if radicate.person.address:
                pdf.multi_cell(0, 5, radicate.person.address+'\n'+radicate.person.email+'\n'+radicate.person.city.name+' '+radicate.person.city.state.name+'\n\n')
            else:
                pdf.multi_cell(0, 5, '\n'+radicate.person.email+'\n'+radicate.person.city.name+' '+radicate.person.city.state.name+'\n\n')
        else:
            pdf.multi_cell(0, 5, radicate.email_user_name+'\n')
            pdf.set_font('Arial', '', 12)
            pdf.multi_cell(0, 5, radicate.email_user_email+'\n\n')
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(20, 5, 'ASUNTO: ')
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0,5, radicate.subject + '\n\n')
        pdf.multi_cell(0,5, radicate.data + '\n\n\n\n\n\nCordialmente,\n\n\n\nRINO GESTI??N DOCUMENTAL'+
                       '\nJos?? Gonz??les\nGerente\n\n')
        pdf.set_font('Arial', '', 7)
        
        annexes = ''
        for index, val in enumerate(list(radicate.files.all())):
            annexes += val.name
            if index != len(list(radicate.files.all())) - 1:
                annexes += ', '
        
        pdf.multi_cell(0, 5, 'Anexo (s): ' + annexes)
        
        # Create custom doc footer
        pdf.custom_footer()
        # Define total number of pages
        pdf.alias_nb_pages()
        # Save pdf file
        # return pdf
        #pdf.output('answer.pdf', 'F')
        return PdfCreationService._save_pdf_in_ecm(pdf, radicate, 'answer' + radicate.number, '.pdf')

class DocxCreationService(object):
    
    @classmethod
    def _save_doc_in_ecm(cls, doc, pqrs, name, extension):
        # Create in-memory buffer
        file_stream = io.BytesIO(doc)
        # Reset the buffer's file-pointer to the beginning of the file
        size = file_stream.tell()
        file_stream.seek(0)
        file_stream.flush()

        node_id = ECMService.upload(File(file_stream, name=name+extension), pqrs.folder_id)

        alfrescoFile = AlfrescoFile(cmis_id=node_id, 
                                    radicate=pqrs,
                                    name=name,
                                    extension=extension,
                                    size=int(size))
        alfrescoFile.save()
        return node_id

    @classmethod
    def _process_doc_to_pdf(cls, doc):
        ###TODO this should be loaded from memory
        doc.save(os.path.join(settings.BASE_DIR, 'media/output.docx'))

        files = {'files': open(os.path.join(
            settings.BASE_DIR, 'media/output.docx'), 'rb')}
        response = requests.post(
            settings.CONVERT_URL,
            files=files
        )
        return response.content

    @classmethod
    def _replace_paragraph(cls, paragraph, pqrs):
        inline = paragraph.runs
        for i in range(len(inline)):
            if len(inline)>i+2 and '<param>' in inline[i].text and '</param>' in inline[i+2].text:
                text_from = inline[i].text + inline[i+1].text + inline[i+2].text
                inline[i].text = FormatHelper.replace_data(text_from, pqrs)
                inline[i+1].text = ''
                inline[i+2].text = ''
            elif len(inline)>i+1 and '<param>' in inline[i].text and '</param>' in inline[i+1].text:
                text_from = inline[i].text + inline[i+1].text
                inline[i].text = FormatHelper.replace_data(text_from, pqrs)
                inline[i+1].text = ''
            elif '<param>' in inline[i].text and '</param>' in inline[i].text:
                inline[i].text = FormatHelper.replace_data(inline[i].text, pqrs)

    @classmethod
    def _process_section(cls, section, pqrs):
        for p in section.paragraphs:
            cls._replace_paragraph(p, pqrs)
        
        for table in section.tables:
            for i, row in enumerate(table.rows):
                for c in row.cells:
                    for p in c.paragraphs:
                        cls._replace_paragraph(p, pqrs)

    @classmethod
    def mix_from_template(cls, template_type, pqrs):
        try:
            template = Template.objects.get(type=template_type)
            doc = Document(template.file)
            cls._process_section(doc, pqrs)

            for section in doc.sections:
                header = section.header
                cls._process_section(header, pqrs)
            
            doc = cls._process_doc_to_pdf(doc)
            cls._save_doc_in_ecm(doc, pqrs, pqrs.number, '.pdf')

        except Template.DoesNotExist:
            template = None

class HtmlPdfCreationService(object):
    @classmethod
    def _save_pdf_in_ecm(cls, file_stream, pqrs, name, extension):
        size = file_stream.tell()
        file_stream.seek(0)
        file_stream.flush()

        node_id = ECMService.upload(File(file_stream, name=name+extension), pqrs.folder_id)

        alfrescoFile = AlfrescoFile(cmis_id=node_id, 
                                    radicate=pqrs,
                                    name=name,
                                    extension=extension,
                                    size=int(size))
        alfrescoFile.save()
        return node_id

    @classmethod
    def generate_pdf(cls, template_type, body, pqrs):
        output_name = ''
        try:
            template = Template.objects.get(type=template_type)
            options = {
                '--load-error-handling': 'skip',
            }
            with NamedTemporaryFile(suffix='.pdf', delete=False) as output:
                header_text = FormatHelper.replace_data(template.header_file.read().decode("utf-8"), pqrs)
                with NamedTemporaryFile(mode='w', suffix='.html', delete=False) as header_html:
                    header_html.write(header_text)
                    options['header-html'] = header_html.name
                
                footer_text = FormatHelper.replace_data(template.footer_file.read().decode("utf-8"), pqrs)
                with NamedTemporaryFile(mode='w', suffix='.html', delete=False) as footer_html:
                    footer_html.write(footer_text)
                    options['footer-html'] = footer_html.name

                output_name = output.name
                pdfkit.from_string(body, output_name, options=options)
                return HtmlPdfCreationService._save_pdf_in_ecm(output, pqrs, pqrs.number, '.pdf')
        finally:
            os.remove(output_name)
            os.remove(options['header-html'])
            os.remove(options['footer-html'])

class Scheduler(object):

    _running = False

    _tasks = {
        'TASK1': lambda : Scheduler.task1(),
        'TASK2': lambda : Scheduler.task2()
    }

    @staticmethod
    def task1():
        print('Executing Task 1...')
        time.sleep(2)
        print('Task 1, done')

    @staticmethod
    def task2():
        print('Executing Task 2...')
        time.sleep(2)
        print('Task 2, done')
        #consecutive = RecordCodeService.get_proceedings_consecutive()
        #print(consecutive)


    @classmethod
    def worker(cls, time_delay):
        while True:
            time.sleep(time_delay)
            for task in Task.objects.all():
                if task.code in cls._tasks:
                    if cls._tasks[task.code]:
                        if cls._check(task):
                            cls._eval(task, datetime.now()) 
                    else:
                        logger.error(f"Task '{task.code}' unimplemented")
                else:
                    logger.error(f"Task '{task.code}' undefined")

    @classmethod
    def _check(cls, task):
        periodicity = task.periodicity
        last = task.last_execution_time
        now = datetime.now()
        if task.status == 1:
            return False
        if not last:
            return True
        if periodicity.isnumeric():
            difference = (now - last).total_seconds()
            return difference >= int(periodicity)
        iter = croniter(periodicity, last)
        return now >= iter.get_next(datetime)
        

    @classmethod
    def _eval(cls, task, now):

        with transaction.atomic():
            try:
                task = Task.objects.filter(id=task.id).select_for_update(nowait=True).get()
                if task.status == 0:
                    task.status = 1
                    task.save()
            except OperationalError as e:
                logger.error(f"Task '{task.code}' status updating")
                return

        try:
            cls._tasks[task.code]()
        except Exception as Error:
            logger.error(f'Task error: {Error}')

        task.last_execution_time = now
        task.status = 0 
        task.save()


    @classmethod
    def start(cls):
        if cls._running:
            return 
        cls._running = True
        t = threading.Thread(target=cls.worker, args=(5,))
        t.setDaemon(True)
        t.start()