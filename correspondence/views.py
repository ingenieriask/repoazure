from pqrs.views import _create_record, _process_next_action
from django import http
from correspondence.models import AlfrescoFile, Radicate, RadicateTypes, ReceptionMode, Record, PermissionRelationAssignation, PermissionRelationReport, ProcessActionStep, \
    RequestInternalInfo
from core.models import Attorny, AttornyType, ChatRooms, FunctionalArea, NotificationsService, Person, Atttorny_Person, Template, UserProfileInfo, FunctionalAreaUser, Alert
from pqrs.models import InterestGroup, PQRS, PqrsContent, SubType, Type
from correspondence.forms import CorrespondenceRadicateForm, RadicateForm, SearchForm, SearchUserForm, UserForm, UserProfileInfoForm, PersonForm, RecordForm, \
    SearchContentForm, ChangeCurrentUserForm, ChangeRecordAssignedForm, LoginForm, AssignToUserForm, ReturnToLastUserForm, ReportToUserForm, \
    DeleteFromReportedForm, RequestInternalInformatioForm
from pqrs.forms import PersonForm as PqrsPeronForm 
from pqrs.forms import PersonAttorny as PqrsPeronAttornyForm 
from pqrs.forms import PersonFormUpdate as PqrsPeronUpdateForm 
from pqrs.forms import ChangeClassificationForm as PqrsChangueType
from django.contrib.auth.models import User
from datetime import datetime
from django.conf import settings
from django.shortcuts import get_list_or_404, get_object_or_404, redirect, render
from django.http import HttpResponseRedirect, HttpResponse, JsonResponse
from django.urls import reverse
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.views.generic.detail import DetailView
from django.views.generic.edit import CreateView
from django.views.generic.list import ListView
from django.views.generic.edit import UpdateView
from django.views.generic import View
from django.contrib.postgres.search import SearchVector, SearchQuery
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q
from django.contrib import messages
from django.core.files.storage import FileSystemStorage, default_storage
from django.core.files.temp import NamedTemporaryFile
from django.contrib.auth.models import Permission
from django.views.decorators.clickjacking import xframe_options_exempt
from django.utils.http import urlencode
from django.core.exceptions import ValidationError
import requests
import json
import os
import io
from docx import Document
import logging
import xlsxwriter
from pinax.eventlog.models import log, Log
from crum import get_current_user
from django.core.files import File
from pydocx import PyDocX

from core.services import DocxCreationService, NotificationsHandler, Recipients, RecordCodeService, SystemParameterHelper, UserHelper
from correspondence.services import ECMService, RadicateService

logger = logging.getLogger(__name__)

# Index view


def index(request):
    return render(request, 'correspondence/index.html', {})


def register(request):
    registered = False
    message = None

    if request.method == 'POST':

        user_form = UserForm(data=request.POST)
        user_profile_form = UserProfileInfoForm(data=request.POST)

        if user_form.is_valid() and user_profile_form.is_valid():
            user = user_form.save()
            user.set_password(user.password)
            user.save()
            profile = user_profile_form.save(commit=False)
            profile.user = user

            if 'picture' in request.FILES:
                profile.picture = request.FILES['picture']

                profile.save()
                registered = True

            message = "El usuario ha sido creado con éxito"
            user_form = UserForm()
        else:
            logger.error(user_form.errors, user_profile_form.errors)
    else:
        user_form = UserForm()
        user_profile_form = UserProfileInfoForm()

    return render(request, 'correspondence/registration.html',
                  context={'user_form': user_form, 'user_profile_form': user_profile_form, 'message': message})


@login_required
def user_logout(request):
    logout(request)
    return HttpResponseRedirect(reverse('correspondence:index'))


# Search by content

@login_required
def search_by_content(request):
    term = ''
    form = None
    radicates = None
    radicate_list = []

    if request.method == 'GET':
        term = request.GET.get('term')
        if not term:
            form = SearchContentForm()
            return render(request, "correspondence/content_search.html", context={'radicates': radicates, 'form': form, 'term': term})
        else:
            form = SearchContentForm(data=request.GET)

    try:
        radicate_list = ECMService.search_by_term(term)
    except (requests.exceptions.Timeout, requests.exceptions.ConnectionError) as Error:
        messages.error(
            request, f"No se ha establecido la conexión con el gestor de contenido")

    if not radicate_list or not radicate_list.count():
        messages.info(
            request, "No se ha encontrado el término en el contenido")

    paginator = Paginator(radicate_list, 10)
    page = request.GET.get('page')

    try:
        radicates = paginator.page(page)
    except PageNotAnInteger:
        radicates = paginator.page(1)
    except EmptyPage:
        radicates = paginator.page(paginator.num_pages)

    return render(request, "correspondence/content_search.html", context={'radicates': radicates, 'form': form, 'term': term})


# Search by names
@login_required
def search_names(request):
    if request.method == 'POST':
        form = SearchForm(request.POST)
        if form.is_valid():
            item = form.cleaned_data['item']
            qs = Person.objects.annotate(
                search=SearchVector('document_number', 'email', 'name', 'address', 'parent__name'), ).filter(
                search=item)
            if not qs.count():
                messages.warning(request, "La búsqueda no obtuvo resultados")

            person_form = PersonForm()
    else:
        form = SearchForm()
        qs = None
        person_form = None

    return render(request, 'correspondence/search.html', context={'form': form, 'list': qs, 'person_form': person_form})


def return_to_last_user(request, radicate):

    pqrs = PqrsContent.objects.get(pk=radicate)
    if request.method == 'POST':
        form = ReturnToLastUserForm(request.POST)
        if form.is_valid():
            user = pqrs.last_user
            pqrs.last_user = pqrs.current_user
            pqrs.current_user = user
            pqrs.pqrsobject.status = PQRS.Status.RETURNED

            if pqrs.observation == None:
                pqrs.observation = ''
            pqrs.observation = pqrs.observation + \
                form.cleaned_data['observations']

            action = ProcessActionStep()
            action.user = get_current_user()
            action.action = 'Devolución'
            action.detail = "El radicado %s ha sido retornado a %s" % (pqrs.number, user.username)
            action.radicate = pqrs
            action.observation = form.cleaned_data['observations']
            action.save()

            log(
                user=request.user,
                action="PQR_RETURNED",
                obj=action,
                extra={
                    "number": pqrs.number,
                    "message": "El radicado %s ha sido retornado a %s" % (pqrs.number, user.username)
                }
            )
            pqrs.save()
            pqrs.pqrsobject.save()
            get_args_str = urlencode({'pk': pqrs.pk, 'template': 'FINISH_RETURN_TO_LAST_USER', 'destination': 'pqrs:radicate_my_inbox'})
            return HttpResponseRedirect(reverse('pqrs:conclusion')+'?'+get_args_str)

    if request.method == 'GET':
        rino_search_user_param = SystemParameterHelper.get(
            'RINO_CORRESPONDENCE_RETURN_TO_LAST_USER').value
        form = ReturnToLastUserForm(request.GET)

    return render(
        request,
        'correspondence/return_to_last_user.html',
        context={
            'form': form,
            'rino_parameter': rino_search_user_param,
            'last_user': pqrs.last_user.username + ' ' + pqrs.last_user.first_name + ' ' + pqrs.last_user.last_name,
            'radicate': radicate
        })


def assign_user(request, radicate):

    if request.method == 'POST':
        form = AssignToUserForm(request.POST)
        if form.is_valid():
            userPk = request.POST.get('selectedUsersInput')
            pqrs = PqrsContent.objects.get(pk=radicate)
            user = User.objects.get(pk=userPk)
            RadicateService.assign_to_user_service(pqrs, user, form.cleaned_data['observations'], 
                reverse('pqrs:detail_pqr', kwargs={'pk': pqrs.pk}), get_current_user(), PQRS.Status.ASSIGNED)
            if pqrs.observation == None:
                pqrs.observation = ''
            pqrs.observation = pqrs.observation + form.cleaned_data['observations']
            
            pqrs.pqrsobject.save()
            pqrs.stage = Radicate.Stage.IN_PROCESS
            pqrs.save()

            get_args_str = urlencode({'pk': pqrs.pk, 'template': 'FINISH_ASSIGNATION', 'destination': 'pqrs:radicate_my_inbox'})
            return HttpResponseRedirect(reverse('pqrs:conclusion')+'?'+get_args_str)
    if request.method == 'GET':
        rino_search_user_param = SystemParameterHelper.get('RINO_CORRESPONDENCE_SEARCH_USER').value
        functional_tree = []
        for item, info in FunctionalArea.get_annotated_list():
            temp = False
            if info['level'] != 0 and int(item.parent.get_depth()+info['level']) > item.get_depth():
                temp = True
            functional_tree.append((item, info, temp))
        form = AssignToUserForm(request.GET)

    return render(
        request,
        'correspondence/assign_user.html',
        context={
            'form': form,
            'rino_parameter': rino_search_user_param,
            'functional_tree': functional_tree,
            'radicate': radicate
        })

filters = {
    1: lambda area, user, permission_list: PermissionRelationAssignation.objects. \
            filter(is_current_area=area, current_permission__in=(user.user_permissions.all() | Permission.objects.filter(group__user=user))). \
            values_list('destination_permission').distinct(),
    2: lambda area, user, permission_list: PermissionRelationReport.objects. \
            filter(is_current_area=area, current_permission__in=(user.user_permissions.all() | Permission.objects.filter(group__user=user))). \
            values_list('destination_permission').distinct(),
    3: lambda area, user, permission_list: Permission.objects.filter(codename__in=set(permission_list)) #('approbation', 'personal_signature', 'legal_signature'))
}

def users_by_area(request):
    filter_pk = request.GET.get('filter_pk')
    kind_task = request.GET.get('kind_task')
    permission_list = request.GET.getlist('permissions[]')
    ###get area from current user and verify if is the same from parameter
    user = get_current_user()
    area = FunctionalAreaUser.objects.filter(Q(user=user) & Q(functional_area=filter_pk)).first() != None
    ###get destination permissions
    ###kind_task 1 is for assination, else is gonna be report
    permission = filters[int(kind_task)](area, user, permission_list)
    ###get destination users
    users = UserHelper.list_by_permissions(permission)
    if request.is_ajax and request.method == "GET":
        users = [{
            'pk': u.user.pk,
            'username': u.user.username,
            'first_name': u.user.first_name,
            'last_name': u.user.last_name
        } for u in FunctionalAreaUser.objects.filter(Q(functional_area=filter_pk) & Q(user__in=users))]
        return JsonResponse(users, safe=False, status=200)
    return JsonResponse({}, status=400)

def get_alerts(request):
    user = get_current_user()
    if request.is_ajax and request.method == "GET":
        alerts = [{
            'pk': a.pk,
            'icon': a.icon,
            'info': a.info,
            'href': a.href,
            'user_creation': (a.user_creation.username + ' ' + a.user_creation.first_name + ' ' + a.user_creation.last_name) if a.user_creation else None
        } for a in Alert.objects.filter(assigned_user=user, is_active = True)]
        return JsonResponse(alerts, safe=False, status=200)
    return JsonResponse({}, status=400)

def disable_alert(request):
    if request.is_ajax and request.method == "POST":
        pk = request.POST.get('pk')
        alert = Alert.objects.get(pk=pk)
        alert.is_active = False
        alert.save()
        return JsonResponse({}, safe=False, status=200)
    return JsonResponse({}, status=400)


def report_to_user(request, radicate):

    if request.method == 'POST':
        form = ReportToUserForm(request.POST)
        if form.is_valid():
            pqrs = PqrsContent.objects.get(pk=radicate)
            RadicateService.report_to_users_service(pqrs, request.POST.getlist('selectedUsersInput'), 
                form.cleaned_data['observations'], reverse('pqrs:reported_detail_pqr', kwargs={'pk': pqrs.pk}), get_current_user())

            if pqrs.observation == None:
                pqrs.observation = ''
            pqrs.observation = pqrs.observation + form.cleaned_data['observations']
            pqrs.save()

            get_args_str = urlencode({'pk': pqrs.pk, 'template': 'FINISH_REPORT', 'destination': 'pqrs:radicate_my_inbox'})
            return HttpResponseRedirect(reverse('pqrs:conclusion')+'?'+get_args_str)
            # return HttpResponseRedirect(reverse('pqrs:radicate_inbox'))

    if request.method == 'GET':
        rino_search_user_param = SystemParameterHelper.get('RINO_CORRESPONDENCE_SEARCH_USER').value
        functional_tree = []
        for item, info in FunctionalArea.get_annotated_list():
            temp = False
            if info['level'] != 0 and int(item.parent.get_depth()+info['level']) > item.get_depth():
                temp = True
            functional_tree.append((item, info, temp))
        form = ReportToUserForm(request.POST)

        return render(
            request,
            'correspondence/report_to_user.html',
            context={
                'form': form,
                'rino_parameter': rino_search_user_param,
                'functional_tree': functional_tree,
                'radicate': radicate
            })

def delete_from_reported(request, radicate):

    if request.method == 'POST':
        form = DeleteFromReportedForm(request.POST)
        if form.is_valid():
            pqrs = PqrsContent.objects.get(pk=radicate)
            current_user = get_current_user()
            users=''
            destination_users = []
            for user in pqrs.reported_people.all():
                
                if user != current_user:
                    destination_users.append(user)
                    users += user.username + ', '
                    
            pqrs.reported_people.remove(current_user)

            
            action = ProcessActionStep()
            action.user = get_current_user()
            action.action = 'Eliminación de informe'
            action.detail = "El usuario %s ha abandonado el informe del radicado %s" % (pqrs.number, current_user)
            action.radicate = pqrs
            action.observation = form.cleaned_data['observations']
            action.save()
            action.destination_users.set(destination_users)
            action.save()

            log(
                user=request.user,
                action="PQR_DELETE_FROM_REPORTED",
                obj=action,
                extra={
                    "number": pqrs.number,
                    "message": "El usuario %s ha abandonado el informe del radicado %s" % (pqrs.number, current_user)
                }
            )
            pqrs.save()
            get_args_str = urlencode({'pk': pqrs.pk, 'template': 'FINISH_DELETE_FROM_REPORTED', 'destination': 'pqrs:radicate_my_inbox'})
            return HttpResponseRedirect(reverse('pqrs:conclusion')+'?'+get_args_str)
            # return HttpResponseRedirect(reverse('pqrs:radicate_inbox'))

    if request.method == 'GET':
        form = DeleteFromReportedForm(request.POST)

        return render(
            request,
            'correspondence/delete_from_reported.html',
            context={
                'form': form,
                'radicate': radicate
            })

# autocomplete


def autocomplete(request):
    if 'term' in request.GET:
        qs = Person.objects.filter(name__icontains=request.GET.get('term'))
        names = list()
        for person in qs:
            names.append(person.name)
        return JsonResponse(names, safe=False)


# Radicate Views
@login_required
def search_user(request):
    form_new=None
    form =SearchUserForm()
    update = False
    if request.method == 'POST':
        form = SearchUserForm(request.POST)
        if form.is_valid():
            check_anonimous = form['anonymous'].value()
            if check_anonimous:
                print(check_anonimous)
            else:
                doc_num = form['doc_num'].value()
                document_type = form['document_type'].value()
                qs = Person.objects.filter(
                    Q(document_number=doc_num) & 
                    Q(document_type=document_type))
                if not qs.count():
                    messages.warning(
                        request, "La búsqueda no obtuvo resultados. Registre la siguiente informacion para continuar con el proceso")
                    form_new = PqrsPeronForm()
                else:
                    person = qs[0]
                    data = {
                        'document_type':person.document_type,
                        'document_number':person.document_number,
                        'expedition_date':person.expedition_date,
                        'name':person.name,
                        'lasts_name':person.lasts_name,
                        'gender_type':person.gender_type,
                        'email':person.email,
                        'email_confirmation':person.email,
                        'request_response':person.request_response,
                        'city':person.city,
                        'phone_number':person.phone_number,
                        'address':person.address,
                        'conflict_victim':person.conflict_victim,
                        'ethnic_group':person.ethnic_group,
                        'preferencial_population':person.preferencial_population.all,
                        'disabilities':person.disabilities.all,
                    }
                    form_new=PqrsPeronUpdateForm(initial=data)
                    update=True
    return render(
        request,
        "correspondence/search_user.html",
        context={
            'form':form,
            'form_new':form_new,
            'update':update
        })


@login_required
def create_pqr_type(request, pqrs):
    pqrsoparent = get_object_or_404(PQRS, uuid=pqrs)
    person = get_object_or_404(Person, id=int(pqrsoparent.principal_person.id))
    form= PqrsChangueType()
    form_new=None
    context=None
    if request.method == 'POST':
        form= PqrsChangueType(request.POST)
        if form.is_valid():
            type = Type.objects.get(pk=form['pqrs_type'].value())
            subtype = SubType.objects.get(pk=form['pqrs_subtype'].value())
            interest_group = InterestGroup.objects.get(pk=form['interest_group'].value())
            form_new = CorrespondenceRadicateForm(
                typePqr=type,
                initial={
                    "subtype_field":subtype,
                    "interestGroup":interest_group
                })
            form_new.person = person
            form= PqrsChangueType()
            context={
                'subtype':subtype,
                'interest_group':interest_group
            }
    return render(
        request,
        "correspondence/create_pqrs.html",
        context={
            'form':form,
            'form_new':form_new,
            'pqrs': pqrs,
            'context':context
        })
    
@login_required
def finish_pqrs(request,pqrs):
    pqrsoparent = get_object_or_404(PQRS, uuid=pqrs)
    person = get_object_or_404(Person, id=int(pqrsoparent.principal_person.id))
    if request.method == 'POST':
        type = SubType.objects.get(pk= request.POST['subtype_field'])
        form = CorrespondenceRadicateForm(type.type, request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.reception_mode = get_object_or_404(ReceptionMode, abbr='VIR')
            instance.type = get_object_or_404(RadicateTypes, abbr='PQR')
            instance.number = RecordCodeService.get_consecutive(RecordCodeService.Type.INPUT)
            instance.response_mode = person.request_response
            instance.person = person
            instance.pqrsobject = pqrsoparent
            radicate =  form.save()
            folder_id = ECMService.create_folder(radicate.number)
            radicate.folder_id = folder_id
            radicate.save()

            action = ProcessActionStep()
            action.user = get_current_user()
            action.action = 'Creación'
            action.detail = 'El radicado %s ha sido creado' % (radicate.number) 
            action.radicate = radicate
            action.save()

            log(
                user=request.user,
                action="PQR_CREATED",
                obj=action,
                extra={
                    "number": radicate.number,
                    "message": "El radicado %s ha sido creado" % (radicate.number)
                }
            )
            query_url = "{0}://{1}/pqrs/consultation/result/{2}".format(request.scheme, request.get_host(), radicate.pk)
            instance.url = query_url
            NotificationsHandler.send_notification('EMAIL_PQR_CREATE', instance,  Recipients(instance.person.email, None, instance.person.phone_number))

            consecutive = 0
            for fileUploaded in request.FILES.getlist('pqrs_creation_uploaded_files'):
                consecutive += 1
                document_temp_file = NamedTemporaryFile()
                for chunk in fileUploaded.chunks():
                    document_temp_file.write(chunk)

                document_temp_file.seek(0)
                document_temp_file.flush()

                node_id = ECMService.upload(File(document_temp_file, name=fileUploaded.name), radicate.folder_id)
                alfrescoFile = AlfrescoFile(cmis_id=node_id, radicate=radicate,
                                            name=os.path.splitext(radicate.number+'-'+str(consecutive).zfill(5))[0],
                                            extension=os.path.splitext(fileUploaded.name)[1],
                                            size=int(fileUploaded.size/1000))
                alfrescoFile.save()

                if not node_id or not ECMService.request_renditions(node_id):
                    messages.error(request, "Ha ocurrido un error al guardar el archivo en el gestor de contenido")

            ### TODO eval in intern creation
            # PdfCreationService.create_pqrs_confirmation_label(radicate)
            DocxCreationService.mix_from_template(Template.Types.PQR_CREATION, instance)
            _process_next_action(instance)
            _create_record(instance)
            messages.success(request, "El radicado se ha creado correctamente")
            url = reverse('pqrs:pqrs_finish_creation', kwargs={'pk': radicate.pk})
            return HttpResponseRedirect(url)
        
        else:
            logger.error("Invalid create pqr form", form.is_valid(), form.errors)
            messages.error( request, "La PQRSD no Pudo ser creada")
    return create_pqr_type(request,pqrs)


@login_required
def create_radicate(request, person):
    from django.core import files

    person = get_object_or_404(Person, id=person)

    if request.method == 'POST':
        form = RadicateForm(request.POST, request.FILES)

        if form.is_valid():
            instance = form.save(commit=False)
            cleaned_data = form.cleaned_data
            form.document_file = request.FILES['document_file']
            now = datetime.now()
            instance.number = now.strftime("%Y%m%d%H%M%S")
            instance.creator = request.user.profile_user
            instance.current_user = request.user.profile_user
            instance.person = person
            radicate = form.save()

            log(
                user=request.user,
                action="RADICATE_CREATED",
                obj=radicate,
                extra={
                    "number": radicate.number,
                    "message": "El radicado %s ha sido creado" % radicate.number
                }
            )
            # TODO make an utility

            # TODO i16n and parameterizable
            NotificationsHandler.send_mail(
                'Notificación RINO: recepción de radicado',
                'Buenos días señor usuario.',
                'rino@skillnet.com.co',
                [instance.person.email]
            )

            document_file = request.FILES['document_file']
            document_temp_file = NamedTemporaryFile()  # ? delete=True)

            for chunk in document_file.chunks():
                document_temp_file.write(chunk)

            document_temp_file.seek(0)
            document_temp_file.flush()

            temp_file = File(document_temp_file, name=document_file.name)

            node_id = ECMService.upload(temp_file)

            if node_id:
                radicate.set_cmis_id(node_id)

                if ECMService.request_renditions(node_id):
                    messages.success(
                        request, "El radicado se ha creado correctamente")
                    url = reverse('correspondence:detail_radicate',
                                  kwargs={'pk': radicate.pk})
                    return HttpResponseRedirect(url)

            messages.error(
                request, "Ha ocurrido un error al guardar el archivo en el gestor de contenido")

        else:
            logger.error("Invalid create radicate form")
            return render(request, 'correspondence/create_radicate.html', context={'form': form, 'person': person})
    else:
        form = RadicateForm(initial={'person': person.id})
        form.person = person

    return render(request, 'correspondence/create_radicate.html', context={'form': form, 'person': person})


class RadicateList(ListView):
    model = Radicate
    context_object_name = 'radicates'

    def get_queryset(self):
        queryset = super(RadicateList, self).get_queryset()
        queryset = queryset.filter(
            current_user=self.request.user.profile_user.pk)
        return queryset


class CurrentUserUpdate(UpdateView):
    model = Radicate
    template_name_suffix = '_currentuser_update_form'
    form_class = ChangeCurrentUserForm


class RecordAssignedUpdate(UpdateView):
    model = Radicate
    template_name_suffix = '_recordassigned_update_form'
    form_class = ChangeRecordAssignedForm

    def form_valid(self, form):

        response = super(RecordAssignedUpdate, self).form_valid(form)

        log(
            user=self.request.user,
            action="RADICATE_ASSIGNED_TO_RECORD",
            obj=self.object,
            extra={
                "number": self.object.number,
                "record": self.object.record.name,
                "message": "El radicado %s ha sido incluído en el expediente %s" % (
                    self.object.number, self.object.record.name)
            }
        )

        if ECMService.assign_record(self.object.cmis_id, self.object.record.cmis_id):
            messages.success(
                self.request, "El archivo se ha guardado correctamente en el expediente")
            return response

        messages.error(
            self.request, "Ha ocurrido un error al actualizar el archivo en el gestor de contenido")
        self.object = None
        return self.form_invalid(form)


def edit_radicate(request, id):
    radicate = get_object_or_404(Radicate, id=id)
    form = RadicateForm(instance=radicate)
    return render(request, 'correspondence/create_radicate.html', context={'form': form, 'person': radicate.person})


class RadicateDetailView(DetailView):
    model = PqrsContent
    template_name="correspondence/radicate_detail.html"
    def get_context_data(self, **kwargs):
        context = super(RadicateDetailView, self).get_context_data(**kwargs)
        context['logs'] = Log.objects.all().filter(object_id=self.kwargs['pk'])
        context['files'] = AlfrescoFile.objects.all().filter(
            radicate=self.kwargs['pk'])
        if context['pqrscontent'].person.attornyCheck:
            personAttorny = Atttorny_Person.objects.filter(
                person=context['pqrscontent'].person.pk)[0]
            context['personAttorny'] = personAttorny
        return context

class RadicateDetailReview(DetailView):
    model = PqrsContent
    template_name="correspondence/radicate_detail.html"
    def get_context_data(self, **kwargs):
        context = super(RadicateDetailReview, self).get_context_data(**kwargs)
        context['logs'] = Log.objects.all().filter(object_id=self.kwargs['pk'])
        context['editable']=True
        context['files'] = AlfrescoFile.objects.all().filter(
            radicate=self.kwargs['pk'])
        if context['pqrscontent'].person.attornyCheck:
            personAttorny = Atttorny_Person.objects.filter(
                person=context['pqrscontent'].person.pk)[0]
            context['personAttorny'] = personAttorny
        return context

def detail_radicate_cmis(request, cmis_id):
    radicate = get_object_or_404(Radicate, cmis_id=cmis_id)
    logs = Log.objects.all().filter(object_id=radicate.pk)
    return render(request, 'correspondence/radicate_detail.html', context={'radicate': radicate, 'logs': logs})


def project_answer(request, pk):
    radicate = get_object_or_404(Radicate, id=pk)
    response_file = None
    answer = ''

    if request.method == 'POST':

        if 'template-file' in request.FILES:
            f = request.FILES['template-file']
            with open(os.path.join(settings.BASE_DIR, settings.MEDIA_DIR, 'temp.docx'), 'wb+') as destination:
                for chunk in f.chunks():
                    destination.write(chunk)
            doc = Document(os.path.join(settings.BASE_DIR, 'media/temp.docx'))

        else:
            doc = Document(os.path.join(settings.BASE_DIR, 'media/template.docx'))

        Dictionary = {
            "*RAD_N*": datetime.now().strftime("%Y%m%d%H%M%S"),
            "*NOMBRES*": str(radicate.person),
            "*CIUDAD*": str(radicate.person.city),
            "*DIRECCION*": str(radicate.person.address) + " - " + str(radicate.person.city),
            "*EMAIL*": str(radicate.person.email),
            "*ASUNTO*": "RESPUESTA " + str(radicate.subject),
            '*FECHA*': datetime.now().strftime("%Y/%m/%d %H:%M:%S"),
            '*ANEXO*': 'Imágenes',
            "*TEXTO*": str(request.POST.get("answer")).replace("\n", ""),
            "*NOMBRES_REMITENTE*": str(radicate.current_user.user.first_name) + " " + str(
                radicate.current_user.user.last_name)
        }

        for line in str(request.POST.get("answer")).replace("\n", ""):
            print(line)

        for i in Dictionary:
            for p in doc.paragraphs:
                if p.text.find(i) >= 0:
                    p.text = p.text.replace(i, Dictionary[i])

        doc.save(os.path.join(settings.BASE_DIR, 'media/output.docx'))

        files = {'files': open(os.path.join(
            settings.BASE_DIR, 'media/output.docx'), 'rb')}

        try:
            response = requests.post(
                settings.CONVERT_URL,
                files=files
            )
            with open(os.path.join(settings.BASE_DIR, 'media/response.pdf'), 'wb') as f:
                f.write(response.content)

            response_file = 'media/response.pdf'
            answer = str(request.POST.get("answer"))

        except Exception as Error:
            logger.error(Error)
            messages.error(request,
                           "Ha ocurrido un error al comunicarse con los servicios de conversión. Por favor informe al administrador del sistema")

    return render(request, 'correspondence/radicate_answer.html',
                  {'radicate': radicate, 'response_file': response_file, 'answer': answer})


# PERSONS Views

class PersonAtronyCreate(CreateView):
    model = Attorny
    form_class = PqrsPeronAttornyForm
    template_name = 'correspondence/person_form.html'

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.save()
        pqrsObject = get_object_or_404(PQRS, uuid=self.kwargs['pqrs_type'])
        attorny_type = get_object_or_404(
            AttornyType, id=int(form['attorny_type'].value()))
        atornyPerson = Atttorny_Person(
            attorny=self.object, person=pqrsObject.principal_person, attorny_type=attorny_type)
        atornyPerson.save()
        return redirect('correspondence:create_pqrs', pqrsObject.uuid)

    def get_form_kwargs(self):
        kwargs = super(PersonAtronyCreate, self).get_form_kwargs()
        # update the kwargs for the form init method with yours
        kwargs.update(self.kwargs)  # self.kwargs contains all url conf params
        return kwargs

class PersonCreateView(CreateView):
    model = Person
    form_class = PqrsPeronForm
    template_name = 'correspondence/person_form.html'

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.save()
        form.save_m2m()
        pqrsObject = PQRS( principal_person=self.object)
        pqrsObject.save()
        if self.object.attornyCheck or form['document_type'].value() == 4:
            return redirect('correspondence:create_person_attorny', pqrsObject.uuid)
        return redirect('correspondence:create_pqrs', pqrsObject.uuid)

class PersonUpdateViewNew(UpdateView):
    model = Person
    form_class = PqrsPeronUpdateForm
    template_name = 'correspondence/person_form.html'

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.save()
        form.save_m2m()
        pqrsObject = PQRS(principal_person=self.object)
        pqrsObject.save()
        if self.object.attornyCheck or form['document_type'].value() == 4:
            return redirect('correspondence:create_person_attorny', pqrsObject.uuid)
        return redirect('correspondence:create_pqrs', pqrsObject.uuid)

class PersonDetailView(DetailView):
    model = Person


class PersonUpdateView(UpdateView):
    model = Person
    form_class = PersonForm


# RECORDS Views
class RecordCreateView(CreateView):
    model = Record
    form_class = RecordForm

    def form_valid(self, form):

        response = super(RecordCreateView, self).form_valid(form)

        id = ECMService.create_record(self.object.name)

        if id:
            self.object.set_cmis_id(id)
            messages.success(
                self.request, "El expediente se ha guardado correctamente")
            return response

        messages.error(
            self.request, "Ha ocurrido un error al crear el expediente en el gestor de contenido")
        self.object = None
        return self.form_invalid(form)


class RecordDetailView(DetailView):
    model = Record


class RecordUpdateView(UpdateView):
    model = Record
    form_class = RecordForm

    def form_valid(self, form):

        response = super(RecordUpdateView, self).form_valid(form)

        if ECMService.update_record(self.object.cmis_id, self.object.name):
            messages.success(
                self.request, "El expediente se ha guardado correctamente")
            return response

        messages.error(
            self.request, "Ha ocurrido un error al actualizar el expediente en el gestor de contenido")
        self.object = None
        return self.form_invalid(form)


class RecordListView(ListView):
    model = Record
    context_object_name = 'records'

    def get_queryset(self):
        queryset = super(RecordListView, self).get_queryset()
        return queryset

    # Charts


def charts(request):
    return render(
        request, 'correspondence/charts.html', 
        context={"chat_rooms":ChatRooms.objects.all(),'current_user':"Usuario Registrado"})


def get_radicates_data(request):
    return Radicate.objects.all().filter(current_user=request.user.profile_user.pk)


class ProcessExcelRadicates(View):

    def get(self, request):
        # Create an in-memory output file for the new workbook.
        output = io.BytesIO()

        # Even though the final file will be in memory the module uses temp
        # files during assembly for efficiency. To avoid this on servers that
        # don't allow temp files, for example the Google APP Engine, set the
        # 'in_memory' Workbook() constructor option as shown in the docs.
        workbook = xlsxwriter.Workbook(output)
        worksheet = workbook.add_worksheet()

        # Add a bold format to use to highlight cells.
        bold = workbook.add_format({'bold': True})

        # Get some data to write to the spreadsheet.
        data = get_radicates_data(request)

        worksheet.set_column(0, 0, 30)
        worksheet.set_column(0, 1, 30)
        worksheet.set_column(0, 2, 30)
        worksheet.set_column(0, 3, 30)
        worksheet.set_column(0, 4, 30)

        worksheet.write('A1', 'Fecha', bold)
        worksheet.write('B1', 'Número', bold)
        worksheet.write('C1', 'Asunto', bold)
        worksheet.write('D1', 'Remitente', bold)
        worksheet.write('E1', 'Estado', bold)

        # Write some test data.
        for row_num, columns in enumerate(data):
            worksheet.write(
                row_num + 1, 0, columns.date_radicated.strftime('%Y-%m-%d'))
            worksheet.write(row_num + 1, 1, columns.number)
            worksheet.write(row_num + 1, 2, columns.subject)
            worksheet.write(row_num + 1, 3, columns.person.name)
            worksheet.write(row_num + 1, 4, columns.type)

        # Close the workbook before sending the data.
        workbook.close()

        # Rewind the buffer.
        output.seek(0)

        # Set up the Http response.
        filename = 'radicados.xlsx'
        response = HttpResponse(
            output,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename=%s' % filename

        return response


@login_required
def get_thumbnail(request):

    cmis_id = request.GET.get('cmis_id')
    prev_response = ECMService.get_thumbnail(cmis_id)
    if prev_response:
        return HttpResponse(prev_response, content_type="image/jpeg")

    return HttpResponse(default_storage.open('tmp/default.jpeg').read(), content_type="image/jpeg")


@xframe_options_exempt
@login_required
def get_file(request):

    cmis_id = request.GET.get('cmis_id')
    prev_response = ECMService.download(cmis_id)
    if prev_response:
        extension = AlfrescoFile.objects.get(cmis_id=cmis_id).extension
        if extension == '.pdf':
            content_type = "application/pdf"
        elif extension == ".doc" or extension == ".docx":
            html = PyDocX.to_html(io.BytesIO(prev_response[0]))
            return HttpResponse(html)
        elif extension == '.jpg':
            content_type = "image/jpeg"
        return HttpResponse(prev_response, content_type=content_type)

    return HttpResponse(default_storage.open('tmp/default.jpeg').read(), content_type="image/jpeg")


def request_internal_info(request, radicate):

    if request.method == 'POST':
        form = RequestInternalInformatioForm(request.POST)
        if form.is_valid():
            instance = RequestInternalInfo(
                assigned_user = User.objects.get(id=request.POST.get('user_selected')),
                description = form.cleaned_data['description'],
                radicate = PqrsContent.objects.get(pk=radicate),
                area = FunctionalArea.objects.get(id=request.POST.get('interest_area'))
            )
            instance.save()
            get_args_str = urlencode({'pk': radicate, 'template': 'FINISH_REQUEST', 
                                      'destination': 'pqrs:radicate_my_inbox',
                                      'status': instance.status,
                                      'addressee': instance.assigned_user})
            return HttpResponseRedirect(reverse('pqrs:conclusion')+'?'+get_args_str)
    if request.method == 'GET':
        functional_tree = []
        for item, info in FunctionalArea.get_annotated_list():
            temp = False
            if info['level'] != 0 and int(item.parent.get_depth()+info['level']) > item.get_depth():
                temp = True
            functional_tree.append((item, info, temp))
        form = RequestInternalInformatioForm(request.GET)

    return render(
        request,
        'correspondence/request_internal_info.html',
        context={
            'form': form,
            'functional_tree': functional_tree,
            'radicate': radicate
        })